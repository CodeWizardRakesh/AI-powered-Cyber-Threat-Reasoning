import os
import pandas as pd
import google.generativeai as genai
from flask import Flask, render_template, request, redirect, url_for, jsonify, Response, send_file
import time
import re
from docx import Document
from io import BytesIO

# Configure Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# **Step 1: Load and Preprocess Data**
def load_and_preprocess_data(file_path):
    """Loads the CSV file and preprocesses it."""
    try:
        df = pd.read_csv(file_path)
        df["Timestamp"] = pd.to_datetime(df["Timestamp"], errors="coerce")
        df.drop_duplicates(inplace=True)
        df.dropna(inplace=True)
        return df
    except Exception as e:
        print(f"Error loading data: {e}")
        return None

# **Step 2: Perform Threat Analysis**
def analyze_threats(df):
    """Analyzes the threat logs for high-risk events."""
    if df is None or isinstance(df, str):
        return None, None
    event_counts = df["Source IP"].value_counts().head(5).to_dict()
    high_risk_events = df[df["Severity"].isin(["High", "Critical"])]
    return event_counts, high_risk_events.to_dict(orient="records")

# **Step 3: Generate LLM-Based Cybersecurity Report**
def generate_threat_report(df):
    """Sends structured logs to Gemini LLM for reasoning and decision-making."""
    if df is None or isinstance(df, str):
        return "Error: No valid data to analyze."
    log_text = df.head(5).to_string(index=False)
    prompt = f"""
    Given the following cybersecurity logs, analyze potential threats and suggest security actions.

    Logs:
    {log_text}

    Instructions:
    1. Summarize key threats in simple terms.
    2. Identify patterns (brute-force attacks, malware, port scans, etc.).
    3. Recommend security actions (e.g., block IP, notify security team, increase firewall rules).

    Respond with fields clearly labeled as:
    "**Summary:** <your summary>"
    "**Threat Patterns:** <your patterns>"
    "**Recommended Actions:** <your actions>"
    """
    try:
        response = model.generate_content(prompt)
        print(f"Raw Gemini response:\n{repr(response.text)}")
        return response.text
    except Exception as e:
        print(f"Error generating report: {e}")
        return "Error: Failed to generate report."

# Generate Word document
def create_word_doc(summary, patterns, actions):
    """Creates a Word document with the report sections."""
    doc = Document()
    doc.add_heading("AI-Generated Cyber Threat Report", 0)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Threat Patterns", level=1)
    doc.add_paragraph(patterns)

    doc.add_heading("Recommended Actions", level=1)
    doc.add_paragraph(actions)

    # Save to a BytesIO object for in-memory handling
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Stream the report with sections separated by delimiters and store for Word doc
def stream_threat_report(df):
    """Streams the threat report word-by-word with section delimiters and prepares Word doc."""
    report = generate_threat_report(df)
    if not report or "Error" in report:
        yield "data: Error generating report.\n\n"
        return

    # Parse the plain text response with Markdown headers
    summary = "No summary provided."
    patterns = "No patterns identified."
    actions = "No actions recommended."

    lines = report.splitlines()
    current_section = None
    for line in lines:
        line = line.strip()
        if re.match(r"\*\*Summary:\*\*.*", line, re.IGNORECASE):
            current_section = "summary"
            summary = line[len("**Summary:**"):].strip()
        elif re.match(r"\*\*Threat Patterns:\*\*.*", line, re.IGNORECASE):
            current_section = "patterns"
            patterns = line[len("**Threat Patterns:**"):].strip()
        elif re.match(r"\*\*Recommended Actions:\*\*.*", line, re.IGNORECASE):
            current_section = "actions"
            actions = line[len("**Recommended Actions:**"):].strip()
        elif current_section and line:
            if current_section == "summary":
                summary += " " + line
            elif current_section == "patterns":
                patterns += " " + line
            elif current_section == "actions":
                actions += " " + line

    if summary == "No summary provided." and patterns == "No patterns identified." and actions == "No actions recommended.":
        print("Warning: No content parsed from Gemini response.")

    # Structure the report with delimiters
    structured_report = (
        "SUMMARY_START " + summary + " SUMMARY_END " +
        "PATTERNS_START " + patterns + " PATTERNS_END " +
        "ACTIONS_START " + actions + " ACTIONS_END"
    )
    print(f"Structured report: {structured_report}")

    # Store the parsed sections in a global variable or session for download
    global report_data
    report_data = {"summary": summary, "patterns": patterns, "actions": actions}

    # Stream word-by-word
    words = structured_report.split()
    for word in words:
        yield f"data: {word}\n\n"
        time.sleep(0.1)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files["file"]
        if file:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(file_path)
            df = load_and_preprocess_data(file_path)
            event_counts, high_risk_events = analyze_threats(df)
            return render_template("index.html", 
                                   event_counts=event_counts,
                                   high_risk_events=high_risk_events,
                                   filename=file.filename)
    return render_template("index.html")

@app.route("/stream-report")
def stream_report():
    filename = request.args.get("filename")
    if not filename:
        return Response("data: No filename provided.\n\n", mimetype="text/event-stream")
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    df = load_and_preprocess_data(file_path)
    return Response(stream_threat_report(df), mimetype="text/event-stream")

@app.route("/download-report")
def download_report():
    """Serves the generated report as a Word document."""
    global report_data
    if not report_data:
        return "No report available for download.", 404

    buffer = create_word_doc(report_data["summary"], report_data["patterns"], report_data["actions"])
    return send_file(
        buffer,
        as_attachment=True,
        download_name="Cyber_Threat_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    report_data = None  # Global variable to store report data
    app.run(debug=True)